"""Document generation service using python-docx."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from app.config import Settings
from app.utils.logger import get_logger

log = get_logger(__name__)


class DocumentService:
    """Generates onboarding documents as .docx files using python-docx."""

    def __init__(self, settings: Settings) -> None:
        self.settings = settings
        self.company_name = settings.company_name

    # ── Public generators ─────────────────────────────────────────────────────

    def generate_offer_letter(self, employee: Any, output_dir: Path) -> Path:
        """Generate a formal offer letter document for the new hire.

        Args:
            employee: Employee ORM model instance.
            output_dir: Directory where the .docx file should be saved.

        Returns:
            Path to the generated .docx file.
        """
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # ── Document styling ──────────────────────────────────────────────────
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # ── Header / company letterhead ───────────────────────────────────────
        header = doc.add_heading(self.company_name, level=0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_heading_color(header, "1a73e8")

        doc.add_paragraph(
            f"Human Resources Department\n"
            f"Date: {datetime.now().strftime('%B %d, %Y')}"
        )
        doc.add_paragraph()

        # ── Recipient ─────────────────────────────────────────────────────────
        doc.add_paragraph(f"Dear {employee.full_name},")
        doc.add_paragraph()

        # ── Body ──────────────────────────────────────────────────────────────
        start_date = str(employee.start_date) if employee.start_date else "your start date"
        job_title = employee.job_title or "Team Member"
        department = employee.department or "the team"

        body = (
            f"We are pleased to offer you the position of {job_title} within "
            f"the {department} department at {self.company_name}. "
            f"Your anticipated start date is {start_date}."
        )
        doc.add_paragraph(body)
        doc.add_paragraph()

        # ── Terms ─────────────────────────────────────────────────────────────
        doc.add_heading("Terms of Employment", level=2)
        terms = [
            ("Position", job_title),
            ("Department", department),
            ("Start Date", start_date),
            ("Employment Type", "Full-Time"),
            ("Location", "As agreed"),
        ]
        table = doc.add_table(rows=len(terms), cols=2)
        table.style = "Table Grid"
        for i, (label, value) in enumerate(terms):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        doc.add_paragraph()

        # ── Acceptance ────────────────────────────────────────────────────────
        doc.add_paragraph(
            "Please confirm your acceptance by signing below. We look forward to "
            "welcoming you to our team."
        )
        doc.add_paragraph()
        doc.add_paragraph(f"Sincerely,\n\n{self.company_name} Human Resources")
        doc.add_paragraph()
        doc.add_paragraph("Employee Signature: ___________________________  Date: __________")

        # ── Save ──────────────────────────────────────────────────────────────
        output_path = output_dir / f"offer_letter_{employee.id}.docx"
        doc.save(str(output_path))
        log.info("Offer letter generated: {path}", path=output_path)
        return output_path

    def generate_it_setup_guide(self, employee: Any, output_dir: Path) -> Path:
        """Generate an IT setup guide document.

        Args:
            employee: Employee ORM model instance.
            output_dir: Directory where the .docx file should be saved.

        Returns:
            Path to the generated .docx file.
        """
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading(f"{self.company_name} — IT Setup Guide", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(
            f"Prepared for: {employee.full_name}\n"
            f"Date: {datetime.now().strftime('%B %d, %Y')}"
        )
        doc.add_paragraph()

        # Sections
        sections = [
            (
                "1. Computer Setup",
                [
                    "Power on your workstation and follow the initial Windows/macOS setup wizard.",
                    "Connect to the corporate Wi-Fi network using your provided credentials.",
                    "Run all pending operating system updates before installing software.",
                    "Install the VPN client (link provided separately by IT).",
                ],
            ),
            (
                "2. Email & Calendar",
                [
                    "Open Outlook or your mail client and sign in with your corporate email.",
                    "Configure your email signature using the company template.",
                    "Accept all calendar invites sent by HR and your manager.",
                    "Enable two-factor authentication (2FA) for your email account.",
                ],
            ),
            (
                "3. Communication Tools",
                [
                    "Install the Slack desktop client and sign in — you've been pre-invited.",
                    "Install Microsoft Teams if applicable to your department.",
                    "Set up your profile picture and status in both tools.",
                ],
            ),
            (
                "4. Jira / Project Tools",
                [
                    "Your Jira account has been provisioned — check your email for the invite.",
                    "Log in at the Jira URL provided and complete profile setup.",
                    "Review any open tickets assigned to you in the Onboarding project.",
                ],
            ),
            (
                "5. Security",
                [
                    "Set a strong, unique password (minimum 12 characters, mix of types).",
                    "Enable full-disk encryption (BitLocker on Windows / FileVault on Mac).",
                    "Complete the mandatory security awareness training within 5 business days.",
                    "Never share credentials or leave your screen unlocked when away.",
                ],
            ),
            (
                "6. IT Support",
                [
                    "For issues, contact the IT Help Desk via the internal ticketing system.",
                    "Emergency contact: it-support@company.com or ext. 1000.",
                ],
            ),
        ]

        for heading, items in sections:
            doc.add_heading(heading, level=2)
            for item in items:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(item)
            doc.add_paragraph()

        output_path = output_dir / f"it_setup_guide_{employee.id}.docx"
        doc.save(str(output_path))
        log.info("IT setup guide generated: {path}", path=output_path)
        return output_path

    def generate_onboarding_checklist(self, employee: Any, output_dir: Path) -> Path:
        """Generate a full onboarding checklist document.

        Args:
            employee: Employee ORM model instance.
            output_dir: Directory where the .docx file should be saved.

        Returns:
            Path to the generated .docx file.
        """
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        title = doc.add_heading(f"{self.company_name} — Onboarding Checklist", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(
            f"Employee: {employee.full_name}\n"
            f"Department: {employee.department or 'N/A'}\n"
            f"Job Title: {employee.job_title or 'N/A'}\n"
            f"Start Date: {employee.start_date or 'TBD'}\n"
            f"Manager: {employee.manager_email or 'N/A'}"
        )
        doc.add_paragraph()

        # Checklist sections
        checklist = [
            (
                "Pre-Start (HR)",
                [
                    ("Send offer letter", "HR"),
                    ("Collect signed contracts", "HR"),
                    ("Set up payroll", "Finance"),
                    ("Order IT equipment", "IT"),
                    ("Provision accounts (email, Slack, Jira)", "IT/HR Automation"),
                ],
            ),
            (
                "Day 1",
                [
                    ("Welcome meeting with HR", "HR"),
                    ("Office/remote workspace tour", "Manager"),
                    ("IT setup completed", "Employee + IT"),
                    ("Complete security awareness training", "Employee"),
                    ("Meet the team", "Manager"),
                ],
            ),
            (
                "Week 1",
                [
                    ("Benefits enrollment", "Employee"),
                    ("Review company policies", "Employee"),
                    ("1:1 with direct manager", "Manager"),
                    ("Set 30-60-90 day goals", "Employee + Manager"),
                    ("Join relevant Slack channels", "Employee"),
                ],
            ),
            (
                "Month 1",
                [
                    ("Complete onboarding training modules", "Employee"),
                    ("30-day check-in with HR", "HR"),
                    ("Submit any outstanding paperwork", "Employee"),
                    ("Performance goals documented", "Employee + Manager"),
                ],
            ),
        ]

        for section_title, items in checklist:
            doc.add_heading(section_title, level=2)
            table = doc.add_table(rows=1 + len(items), cols=3)
            table.style = "Table Grid"

            # Header row
            hdr = table.rows[0].cells
            hdr[0].text = "Task"
            hdr[1].text = "Responsible"
            hdr[2].text = "Status"
            for cell in hdr:
                run = cell.paragraphs[0].runs
                if run:
                    run[0].bold = True

            for i, (task, responsible) in enumerate(items, 1):
                row = table.rows[i].cells
                row[0].text = task
                row[1].text = responsible
                row[2].text = "☐ Pending"

            doc.add_paragraph()

        doc.add_paragraph(
            f"\nDocument generated on {datetime.now().strftime('%Y-%m-%d %H:%M')} "
            f"by {self.company_name} HR Onboarding System."
        )

        output_path = output_dir / f"onboarding_checklist_{employee.id}.docx"
        doc.save(str(output_path))
        log.info("Onboarding checklist generated: {path}", path=output_path)
        return output_path


# ── Helpers ───────────────────────────────────────────────────────────────────


def _set_heading_color(heading, hex_color: str) -> None:
    """Set the font colour of all runs in a docx heading paragraph."""
    try:
        from docx.shared import RGBColor

        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(r, g, b)
    except Exception:
        pass  # Non-critical styling; don't break generation
